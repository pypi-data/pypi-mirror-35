#!/usr/bin/python
# -*- coding: utf-8 -*-
import urllib3
import facebook as fb
import requests
import time
import datetime as dt
from statistics import mean
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlencode
from urllib.request import urlretrieve
import sys
from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from tkinter import messagebox as mb
import pickle
import os
import subprocess
if sys.platform == 'win32':
    sys._enablelegacywindowsfsencoding()

class programo(Tk):
    def __init__(self, parent):
        self.parent = parent
        self.iniciar()

    def iniciar(self):

        def tokenValido(tokendado):
            graph = fb.GraphAPI(access_token = tokendado, version = 3.0)
            try:
                objetoPage = graph.request('me?fields=id,name')
                ingresoNombre()
            except fb.GraphAPIError:
                error = mb.showerror('Error', 'El token ingresado no es válido')
                token.set('')

        def chequearGenerar():
            time = chequearTreinta('a')
            if time is False:
                return
            tokenok = token.get()
            if (len(tokenok) != 0
                and len(criterio.get()) > 0 
                and (len(isfacebook.get()) > 0 
                    or len(isinstagram.get()) > 0)):
                tokenValido(tokenok)
            else:
                error = mb.showerror('Error', 'Falta ingresar datos requeridos')

        def ingresoNombre():
            extension = '.docx'
            futuro = fd.asksaveasfilename(defaultextension = extension,
                filetypes = [("docx", extension)], confirmoverwrite=True)
            if len(futuro) > 0:
                if not futuro.endswith(extension):
                    futuro += extension
                generarInforme(futuro)

        def obtenerMes(numero):
            mes = ('','enero','febrero',
                'marzo','abril','mayo',
                'junio','julio','agosto',
                'septiembre','octubre',
                'noviembre','diciembre')
            return mes[numero]

        def generarInforme(futuro):

            def getfanpageObject():
                objetoPage = graph.request('me?fields=id,name')
                return objetoPage

            def getinstaID(pageid):
                objetoPage = graph.request('/' + pageid + 
                    '?fields=instagram_business_account')
                instaid = objetoPage['instagram_business_account']['id']
                return instaid

            def replaceTime(key, metrica):
                temp = key[metrica].replace("T", " ")
                temp, _ = temp.split('+')
                temp = dt.datetime.strptime(temp, "%Y-%m-%d %H:%M:%S")
                return temp

            def acordePeriodo(insta_or_face, identificadores, inicio, fin):
                id_en_periodo = []
                if insta_or_face is instagram:
                    metrica = 'timestamp'
                else:
                    metrica = 'created_time'
                for i in identificadores:
                    post = graph.request('/'+ i + '?fields=' + metrica)
                    post[metrica] = replaceTime(post, metrica)
                    if post[metrica] >= inicio and post[metrica] <= fin:
                        id_en_periodo.append(i)    
                return id_en_periodo

            def getIdentificadores(insta_or_face, inicio, fin):
                #Extrae todos los posts en el periodo indicado
                #TODO requiere revisar el valor "paging" que permite acceso a más posts
                identificadores = []
                if insta_or_face is instagram:
                    metrica = 'media'
                    fanid = instaid
                    todoslosposts = graph.request('/'+ fanid +'/' + metrica)
                else:
                    metrica = 'feed?limit=100'
                    fanid = pageid
                    todoslosposts = graph.request('/'+ fanid +'/' + metrica + 'since=' + inicio + '&until=' + fin)
                datostodos = todoslosposts['data']
                for x in datostodos:
                    identificadores.append(x['id'])
                return identificadores

            def metricaUsuario(insta_or_face, metrica, inicio, fin, periodo):
                #Extrae todos los posts en el periodo indicado
                if insta_or_face is instagram:
                    fanid = instaid
                else:
                    fanid = pageid
                valores = []
                todoslosposts = graph.request('/'+ fanid + '/insights?metric=' + metrica + 
                    '&period=' + periodo + '&since=' + inicio + '&until=' + fin)
                datostodos = todoslosposts['data']
                for x in datostodos[0]['values']:
                    valores.append(x['value'])
                return valores

            #Extrae valor de la métrica de cada uno de los posts del periodo seleccionado
            def metricaUnica(metrica, identificadores, una_o_dos):
                top = []
                top_metrica = []
                for i in identificadores:
                    metricapost = graph.request('/'+ i +'/insights?metric='+ metrica)
                    datospost = metricapost['data']
                    for x in datospost:
                        u = x['values'][0]['value']
                        top.append(u)
                        top_metrica.append((i, u))
                if una_o_dos is True:
                    return top, top_metrica
                else:
                    return top

            def metricaSumada(valores):
                if type(valores) is dict:
                    u = []
                    for x in valores.values():
                        u.append(x)
                    return sum(u)
                else:
                    return sum(valores)

            def metricasPostPopular(insta_or_face, postpopular):
                dicmetpop = {}
                if insta_or_face is instagram:
                    if postpopular['media_type'] == 'IMAGE':
                        popmetricas = 'impressions,reach,engagement,saved'
                    elif postpopular['media_type'] == 'VIDEO':
                        popmetricas = 'impressions,reach,engagement,saved,video_views'
                else:
                    if postpopular['type'] == 'video':
                        popmetricas = 'post_impressions,post_impressions_unique,post_engaged_users,post_reactions_by_type_total,post_video_avg_time_watched'
                    else:
                        popmetricas = 'post_impressions,post_impressions_unique,post_engaged_users,post_reactions_by_type_total'

                poppostmetricas = graph.request('/' + postpopular['id'] + '/insights?metric=' + popmetricas)
                accesodata = poppostmetricas['data']
                dicmetpop['impresiones'] = accesodata[0]['values'][0]['value']
                dicmetpop['alcance'] = accesodata[1]['values'][0]['value']
                dicmetpop['engagement'] = accesodata[2]['values'][0]['value']
                dicmetpop['guardado'] = accesodata[3]['values'][0]['value']
                if len(accesodata) == 5:
                    dicmetpop['vistas_video'] = accesodata[4]['values'][0]['value']

                return dicmetpop

            #TODO Extraer post popular a partir del engagement rate
            def postPopular(insta_or_face, metrica, identificadores):
                top, top_metrica = metricaUnica(metrica, identificadores, True)
                if insta_or_face is instagram:
                    campos = 'fields=id,media_type,permalink'
                else:
                    campos = 'fields=id,type,permalink_url'

                #Extrae información del post más popular
                for a, b in top_metrica:
                    if b == max(top):
                        postpopular = graph.request('/'+ a + '?' + campos)

                return postpopular

            def getPermalink(insta_or_face, infopost):
                if insta_or_face is instagram:
                    a = infopost['permalink']
                    #Por alguna razón no puedo rescatar la imagen
                    # al usar https.
                    if 'https' in a:
                        a = a.replace('https', 'http')
                    return a
                else:
                    return infopost['permalink_url']

            def escribirDocumentoAmbos(red, *data):

                diccionario = [x for x in data]
                
                if len(diccionario) is 1:
                    if len(diccionario[0]) is 12:
                        facedic = diccionario[0]
                        
                    elif len(diccionario[0]) is 11:
                        instadic = diccionario[0]
                        
                elif len(diccionario) is 2:
                    if len(diccionario[0]) is 12:
                        facedic = diccionario[0]
                        instadic = diccionario[1]
                        
                    elif len(diccionario[0]) is 11:
                        instadic = diccionario[0]
                        facedic = diccionario[1]
                        

                #Nombres de los archivos guardados
                desktop = os.getcwd()
                fscreenshot = os.path.join(desktop, 'screenshotfacebookpost.png')
                iscreenshot = os.path.join(desktop, 'screenshotinstagrampost.png')

                document = Document()
                p = document.add_paragraph()
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run('Informe {0} {1} {2}'.format(pagename, mes, anho)).bold = True

                if (red == 'facebook') or (red == 'ambos'):
                    #Descarga de Facebook Diccionario
                    ftotaldeposts = facedic['totaldeposts']
                    fseguidorestotal = facedic['seguidorestotal']
                    fimpresionespromedio = facedic['impresionespromedio']
                    freachpromedio = facedic['reachpromedio'] 
                    fsavedpromedio = facedic['savedpromedio']
                    fimpresionespop = facedic['impresionespop']
                    fengagementpop = facedic['engagementpop']
                    freachpop = facedic['reachpop']
                    fsavedpop = facedic['savedpop']
                    flinkpost = facedic['linkpost']
                    fnuevolikes = facedic['nuevolikes']
                    #Cuenta de Facebook
                    p2 = document.add_paragraph()
                    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.add_run('Facebook').bold = True

                    p3 = document.add_paragraph()
                    p3.add_run('{0} nuevos likes (llegando a {1}).'.format(fnuevolikes, fseguidorestotal))

                    p4 = document.add_paragraph()
                    p4.add_run('{0} posts emitidos.'.format(ftotaldeposts))

                    p5 = document.add_paragraph()
                    p5.add_run('{0} usuarios en promedio que realizaron algún tipo de interacción.'.format(fsavedpromedio))

                    p6 = document.add_paragraph()
                    p6.add_run('{0} personas alcanzadas (promedio de usuarios únicos en cada post).'.format(freachpromedio))

                    p7 = document.add_paragraph()
                    p7.add_run('{0} impresiones en promedio.'.format(fimpresionespromedio))

                    #TODO tasa de participación como criterio
                    #p75 = document.add_paragraph()
                    #p75.add_run('{0}% tasa de participación (engagement rate) promediada.'.format(mediaER))

                    p8 = document.add_paragraph('Post destacado:')

                    #Acceso a internet y extración de imagen del post más popular.
                    try:
                        paramsf = urlencode({'url':'{0}'.format(flinkpost), 
                            'access_key':'3dfe7afa8bafb96cdb0215bb12974a89'})
                        urlretrieve('https://api.screenshotlayer.com/api/capture?' 
                            + paramsf, fscreenshot)
                        pimage = document.add_picture(fscreenshot, width=Inches(5))
                    except:
                        p00 = document.add_paragraph()
                        p00.add_run('{0}'.format(flinkpost))

                    #p85 = document.add_paragraph()
                    #p85.add_run('{0}% tasa de participación.'.format(masER))

                    p9 = document.add_paragraph()
                    p9.add_run('{0} usuarios únicos alcanzados.'.format(freachpop))

                    p10 = document.add_paragraph()
                    p10.add_run('{0} impresiones.'.format(fimpresionespop))

                    p11 = document.add_paragraph()
                    p11.add_run('{0} usuarios que interactuaron.'.format(fengagementpop))

                    document.add_page_break()

                if (red == 'instagram') or (red == 'ambos'):
                    #Descarga de Instagram Diccionario
                    itotaldeposts = instadic['totaldeposts']
                    iseguidorestotal = instadic['seguidorestotal']
                    iimpresionespromedio = instadic['impresionespromedio']
                    ireachpromedio = instadic['reachpromedio'] 
                    isavedpromedio = instadic['savedpromedio']
                    iimpresionespop = instadic['impresionespop']
                    iengagementpop = instadic['engagementpop']
                    ireachpop = instadic['reachpop']
                    isavedpop = instadic['savedpop']
                    ilinkpost = instadic['linkpost']

                    pgram = document.add_paragraph()
                    pgram.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    pgram.add_run('Instagram').bold = True

                    pgram0 = document.add_paragraph()
                    pgram0.add_run('{0} nuevos seguidores.'.format(iseguidorestotal))

                    pgram1 = document.add_paragraph()
                    pgram1.add_run('{0} posts realizados.'.format(itotaldeposts))

                    pgram2 = document.add_paragraph()
                    pgram2.add_run('{0} impresiones (en promedio).'.format(iimpresionespromedio))

                    pgram3 = document.add_paragraph()
                    pgram3.add_run('{0} usuarios alcanzados (en promedio).'.format(ireachpromedio))

                    pgram4 = document.add_paragraph()
                    pgram4.add_run('{0} posts guardados.'.format(isavedpromedio))

                    pgram5 = document.add_paragraph('Post destacado:')

                    #Acceso a internet y extración de imagen del post más popular.
                    try:
                        paramsi = urlencode({'url':'{0}'.format(ilinkpost), 
                            'access_key':'3dfe7afa8bafb96cdb0215bb12974a89'})
                        urlretrieve('https://api.screenshotlayer.com/api/capture?' 
                            + paramsi, iscreenshot)
                        pimage2 = document.add_picture(iscreenshot, width=Inches(5))
                    except:
                        pgram00 = document.add_paragraph()
                        pgram00.add_run('{0}'.format(ilinkpost))

                    pgram6 = document.add_paragraph()
                    pgram6.add_run('{0} interacciones de usuarios (likes + comentarios).'.format(iengagementpop))

                    pgram7 = document.add_paragraph()
                    pgram7.add_run('{0} impresiones.'.format(iimpresionespop))

                    pgram8 = document.add_paragraph()
                    pgram8.add_run('{0} usuarios alcanzados.'.format(ireachpop))

                    document.add_page_break()

                #Guardado del documento
                document.save('{0}'.format(futuro))

                #Eliminación de imágenes
                if os.path.isfile(fscreenshot):
                    os.remove(fscreenshot)

                if os.path.isfile(iscreenshot):
                    os.remove(iscreenshot)

                #Apertura de archivo
                if sys.platform == 'linux':
                    os.system("xdg-open " + futuro)
                elif sys.platform == 'win32':
                    os.startfile(futuro)
                else:
                    opener = "open" if sys.platform == 'darwin' else 'xdg-open'
                    subprocess.call([opener, futuro])

            def getCriterio(insta_or_face):
                eleccion = criterio.get()
                if insta_or_face is instagram:
                    if eleccion == 'impressions':
                        return 'impressions'
                    elif eleccion == 'reach':
                        return 'reach'
                    elif eleccion == 'engagement':
                        return 'engagement'
                else:
                    if eleccion == 'impressions':
                        return 'post_impressions'
                    elif eleccion == 'reach':
                        return 'post_impressions_unique'
                    elif eleccion == 'engagement':
                        return 'post_engaged_users'


            def getDatos(insta_or_face):
                if insta_or_face is instagram:
                    c_followers = 'follower_count'
                    c_impresiones = 'impressions'
                    c_alcance = 'reach'
                    c_guardado = 'saved'
                    c_periodo = 'day'
                    eleccionpop = getCriterio(instagram)
                else:
                    c_followers = 'page_fans'
                    c_impresiones = 'post_impressions'
                    c_alcance = 'post_impressions_unique'
                    c_guardado = 'post_engaged_users'
                    c_periodo = 'lifetime'
                    eleccionpop = getCriterio(facebook)

                #Obtención de datos
                identificadores = getIdentificadores(insta_or_face, iniciounix, finunix)
                identificadores = acordePeriodo(insta_or_face, identificadores, inicito, finito)
                totaldeposts = len(identificadores)

                seguidores = metricaUsuario(insta_or_face, c_followers, iniciounix, finunix, c_periodo)
                if insta_or_face is instagram:
                    seguidorestotal = metricaSumada(seguidores)
                else:
                    seguidorestotal = seguidores[-1]

                valoresimpresiones = metricaUnica(c_impresiones, identificadores, False)
                impresionespromedio = mean(valoresimpresiones) if len(valoresimpresiones) > 0 else 0
                valoresreach = metricaUnica(c_alcance, identificadores, False)
                reachpromedio = mean(valoresreach) if len(valoresimpresiones) > 0 else 0
                valoressaved = metricaUnica(c_guardado, identificadores, False)
                savedpromedio = mean(valoressaved) if len(valoresimpresiones) > 0 else 0

                infopostpop = postPopular(insta_or_face, eleccionpop, identificadores)
                diccionariopop = metricasPostPopular(insta_or_face, infopostpop)

                impresionespop = diccionariopop['impresiones']
                engagementpop = diccionariopop['engagement']
                reachpop = diccionariopop['alcance']
                savedpop = diccionariopop['guardado']
                
                linkpost = getPermalink(insta_or_face, infopostpop)

                #Edición de datos para display en documento
                impresionespromedio = str(round(impresionespromedio, 1)).replace('.',',')
                reachpromedio = str(round(reachpromedio, 1)).replace('.',',')
                savedpromedio = str(round(savedpromedio, 1)).replace('.', ',')

                todoslosdatos = {'totaldeposts' : totaldeposts, 'seguidorestotal' : seguidorestotal, 
                'impresionespromedio' : impresionespromedio, 'reachpromedio' : reachpromedio, 
                'savedpromedio' : savedpromedio, 'impresionespop' : impresionespop, 
                'engagementpop' : engagementpop, 'reachpop' : reachpop, 'savedpop' : savedpop, 
                'linkpost' : linkpost}

                if 'vistas_video' in diccionariopop:
                    vistaspop = diccionariopop['vistas_video']
                    todoslosdatos['vistaspop'] = vistaspop

                if insta_or_face is facebook:
                    nuevolikes = seguidores[-1] - seguidores[0]
                    todoslosdatos['nuevolikes'] = nuevolikes

                return todoslosdatos

            facebook = False
            instagram = True
            tokenok = token.get()

            #Objeto con acceso a Facebook e Instagram
            graph = fb.GraphAPI(access_token=tokenok, version= 3.0)

            #IDs de páginas
            pageobjeto = getfanpageObject()
            pagename = pageobjeto['name']
            pageid = pageobjeto['id']
            instaid = getinstaID(pageid)

            #Rango temporal
            diainicio = int(dia_inicio.get())
            mesinicio = mesNumero(mes_inicio.get())
            anhoinicio = int(anho_inicio.get())
            diafin = int(dia_fin.get())
            mesfin = mesNumero(mes_fin.get())
            anhofin = int(anho_fin.get())
            inicito = dt.datetime(anhoinicio, mesinicio, diainicio)
            finito = dt.datetime(anhofin, mesfin, diafin)

            #iConversión a Timestamp string
            iniciounix = str(int(time.mktime(inicito.timetuple()))) 
            finunix = str(int(time.mktime(finito.timetuple())))

            #Definir periodo analizado en el informe
            if int(inicito.year) == int(finito.year):
                if int(inicito.month) == int(finito.month):
                    mes = obtenerMes(inicito.month)
                else:
                    mes = (str(int(inicito.day)) + "/" + str(int(inicito.month)) + 
                        " al " + str(int(finito.day)) + "/" 
                        + str(int(finito.month)) + " del ")
                anho = int(finito.year)
            else:
                mes = (str(int(inicito.day)) + "/" + str(int(inicito.month)) + 
                    "/" + str(int(inicito.year))+ " al " + str(int(finito.day)) +
                     "/" + str(int(finito.month)) + "/" + str(int(finito.year)))
                anho = ""

            if (isfacebook.get() == 'facebook' and isinstagram.get() == 'instagram'):
                allthedataface = getDatos(facebook)
                allthedatainsta = getDatos(instagram)
                escribirDocumentoAmbos('ambos', allthedataface, allthedatainsta)
            elif (isfacebook.get() == 'nofacebook' and isinstagram.get() == 'instagram') or (isinstagram.get() == 'instagram'):
                allthedata = getDatos(instagram)
                escribirDocumentoAmbos('instagram', allthedata)
            elif (isfacebook.get() == 'facebook' and isinstagram.get() == 'noinstagram') or (isfacebook.get() == 'facebook'):
                allthedata = getDatos(facebook)
                escribirDocumentoAmbos('facebook', allthedata)

        def habilitarMesInicio(event):
            listamesesi.config(state=['readonly'])

        def habilitarAnhoInicio(event):
            listaanhosi.config(state=['readonly'])

        def habilitarDiaFin(event):
            listadiasf.config(state=['readonly'])

        def habilitarMesFin(event):
            listamesesf.config(state=['readonly'])

        def habilitarAnhoFin(event):
            listaanhosf.config(state=['readonly'])

        def mesNumero(mestexto):
            textonumero = {'enero':1,'febrero':2,
                'marzo':3,'abril':4,'mayo':5,
                'junio':6,'julio':7,'agosto':8,
                'septiembre':9,'octubre':10,
                'noviembre':11,'diciembre':12}
            return textonumero[mestexto] if mestexto in textonumero else 0

        def chequearTreinta(event):
            a = dia_inicio.get() or 0
            b = mes_inicio.get() or 0
            c = anho_inicio.get() or 0
            d = dia_fin.get() or 0
            e = mes_fin.get() or 0
            f = anho_fin.get() or 0

            diainicio = int(a)
            mesinicio = mesNumero(mes_inicio.get())
            anhoinicio = int(c)
            diafin = int(d)
            mesfin = mesNumero(mes_fin.get())
            anhofin = int(f)

            suma = diainicio + mesinicio + anhoinicio + diafin + mesfin + anhofin

            if suma > 4024:
                try:
                    fechainicio = dt.datetime(anhoinicio, mesinicio, diainicio)
                except:
                    if mesinicio is 2:
                        error = mb.showerror('Error', 'El mes de {0} tiene sólo 29 días'.format(mes_inicio.get()))
                        return False
                    else:
                        error = mb.showerror('Error', 'El mes de {0} no tiene 31 días'.format(mes_inicio.get()))
                        return False

                try:
                    fechatermino = dt.datetime(anhofin, mesfin, diafin)
                except:
                    if mesinicio is 2:
                        error = mb.showerror('Error', 'El mes de {0} tiene sólo 29 días'.format(mes_inicio.get()))
                        return False
                    else:
                        error = mb.showerror('Error', 'El mes de {0} no tiene 31 días'.format(mes_fin.get()))
                        return False

                if fechainicio < fechatermino:
                    rangoidealdetermino = fechainicio + dt.timedelta(days=30)
                    if fechatermino <= rangoidealdetermino:
                        return True
                    else:
                        error = mb.showerror('Error', 'El rango entre fechas no debe superar los 30 días.')
                        return False
                else:
                    error = mb.showerror('Error', 'La fecha de inicio debe ser anterior a la fecha de termino.')
                    return False
            else:
                pass


        self.parent.columnconfigure(0, weight = 1)
        self.parent.rowconfigure(0, weight = 1)

        #Año actual para ingresar a la lista
        anhoactual = dt.datetime.today().year

        #Tupla con los meses en string
        mesestexto = ('enero','febrero',
                'marzo','abril','mayo',
                'junio','julio','agosto',
                'septiembre','octubre',
                'noviembre','diciembre')
        
        #Creación de base de datos
        desktop = os.getcwd()

        frame1 = ttk.Frame(self.parent, borderwidth = 1, 
            padding = (4, 4, 4, 4))
        frame1.grid(column = 0, row = 0, sticky = (W, E, S, N))

        frase = ttk.Label(frame1, text = 'Fecha Inicio: ')
        frase.grid(column = 0, row = 1)

        dia_inicio = StringVar()
        listadiasi = ttk.Combobox(frame1, textvariable = dia_inicio)
        listadiasi['values'] = [i for i in range(1, 32)]
        listadiasi.state(['readonly'])
        listadiasi.grid(column = 1, row = 1)
        #listadiasi.bind("<<ComboboxSelected>>", habilitarMesInicio)
        listadiasi.bind("<<ComboboxSelected>>", chequearTreinta)

        mes_inicio = StringVar()
        listamesesi = ttk.Combobox(frame1, textvariable = mes_inicio)
        listamesesi['values'] = mesestexto
        listamesesi.state(['readonly'])
        listamesesi.grid(column = 2, row = 1)
        #listamesesi.bind("<<ComboboxSelected>>", habilitarAnhoInicio)
        listamesesi.bind("<<ComboboxSelected>>", chequearTreinta)

        anho_inicio = StringVar()
        listaanhosi = ttk.Combobox(frame1, textvariable = anho_inicio)
        listaanhosi['values'] = [i for i in range(2010, anhoactual + 1)]
        listaanhosi.state(['readonly'])
        listaanhosi.grid(column = 3, row = 1)
        #listaanhosi.bind("<<ComboboxSelected>>", habilitarDiaFin)
        listaanhosi.bind("<<ComboboxSelected>>", chequearTreinta)

        frase2 = ttk.Label(frame1, text = 'Fecha Termino: ')
        frase2.grid(column = 0, row = 2)

        dia_fin = StringVar()
        listadiasf = ttk.Combobox(frame1, textvariable = dia_fin)
        listadiasf['values'] = [i for i in range(1, 32)]
        listadiasf.state(['readonly'])
        listadiasf.grid(column = 1, row = 2)
        #listadiasf.bind("<<ComboboxSelected>>", habilitarMesFin)
        listadiasf.bind("<<ComboboxSelected>>", chequearTreinta)

        mes_fin = StringVar()
        listamesesf = ttk.Combobox(frame1, textvariable = mes_fin)
        listamesesf['values'] = mesestexto
        listamesesf.state(['readonly'])
        listamesesf.grid(column = 2, row = 2)
        #listamesesf.bind("<<ComboboxSelected>>", habilitarAnhoFin)
        listamesesf.bind("<<ComboboxSelected>>", chequearTreinta)

        anho_fin = StringVar()
        listaanhosf = ttk.Combobox(frame1, textvariable = anho_fin)
        listaanhosf['values'] = [i for i in range(2010, anhoactual + 1)]
        listaanhosf.state(['readonly'])
        listaanhosf.grid(column = 3, row = 2)
        listaanhosf.bind("<<ComboboxSelected>>", chequearTreinta)

        s1 = ttk.Separator(frame1, orient = HORIZONTAL)
        s1.grid(row = 3, columnspan = 8, sticky = (W, E))

        seleccionPop = ttk.Label(frame1, 
            text = 'Seleccionar criterio para\n clasificar posts')
        seleccionPop.grid(column = 0, row = 4, sticky = (N), pady = 2)

        criterio = StringVar(None, '')
        por_impresion = ttk.Radiobutton(frame1, text = 'Impresiones', 
            variable = criterio, value = 'impressions')
        por_alcance = ttk.Radiobutton(frame1, text = 'Alcance', 
            variable = criterio, value = 'reach')
        #por_engagement = ttk.Radiobutton(frame1, text = 'Tasa de participación', 
        #    variable = criterio, value = 'engagement_rate')
        por_usuarios = ttk.Radiobutton(frame1, text = 'Acciones de usuarios', 
            variable = criterio, value = 'engagement')
        
        por_impresion.grid(column = 0, row = 5, sticky = W)
        por_alcance.grid(column = 0, row = 6, sticky = W)
        #por_engagement.grid(column = 0, row = 8, sticky = W)
        por_usuarios.grid(column = 0, row = 7, sticky = W)

        s2 = ttk.Separator(frame1, orient=VERTICAL)
        s2.grid(column = 1, row = 3, rowspan = 5, sticky = (N, S))

        ingreso_archivos = ttk.Label(frame1, text = 'Ingrese token')
        ingreso_archivos.grid(column = 2, row = 4, sticky = (E, W), 
            columnspan = 2, pady = 2)

        token = StringVar()
        entrada_post = ttk.Entry(frame1, textvariable = token)
        entrada_post.grid(column = 2, row = 5, sticky = (E, W))

        s4 = ttk.Separator(frame1, orient=VERTICAL)
        s4.grid(column = 3, row = 3, rowspan = 5, sticky = (N, S))

        seleccionredes = ttk.Label(frame1, text = 'Seleccione redes sociales')
        seleccionredes.grid(column = 4, row = 4, pady=2, columnspan = 4)

        isfacebook = StringVar()
        facebookcheck = ttk.Checkbutton(frame1, text='Facebook', 
            variable=isfacebook, onvalue='facebook', offvalue='nofacebook')
        facebookcheck.grid(column = 4, row = 5, sticky = (E, W),
            columnspan = 2, pady = 2)

        isinstagram = StringVar()
        instagramcheck = ttk.Checkbutton(frame1, text='Instagram',
            variable=isinstagram, onvalue='instagram', offvalue= 'noinstagram')
        instagramcheck.grid(column = 4, row = 6, sticky = (E, W),
            columnspan = 2, pady = 2)

        s3 = ttk.Separator(frame1, orient = HORIZONTAL)
        s3.grid(row = 9, column = 0, columnspan = 8, sticky = (E, W))

        generar = ttk.Button(frame1, text = 'Generar Informe', 
            command = chequearGenerar)
        generar.grid(row = 10, column = 0, columnspan = 5)

def main():
    root = Tk()
    root.title('Informe de Instagram')
    app = programo(root)
    root.mainloop()

if __name__ == "__main__":
    main()