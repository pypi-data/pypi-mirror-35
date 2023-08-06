#!/usr/bin/python
# -*- coding: utf-8 -*-
import urllib3
import facebook
import requests
import time
import datetime as dt
from statistics import mean
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlencode
from urllib.request import urlretrieve
import sys
from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from tkinter import messagebox as mb
import pickle
import os
import subprocess
if sys.platform == 'win32':
    sys._enablelegacywindowsfsencoding()

class programo(Tk):
    def __init__(self, parent):
        self.parent = parent
        self.iniciar()

    def iniciar(self):

        def tokenValido(tokendado):
            graph = facebook.GraphAPI(access_token = tokendado, version = 3.0)
            try:
                objetoPage = graph.request('me?fields=id,name')
                ingresoNombre()
            except facebook.GraphAPIError:
                error = mb.showerror('Error', 'El token ingresado no es válido')
                token.set('')

        def chequearGenerar():
            tokenok = token.get()
            if (len(tokenok) != 0
                and len(criterio.get()) > 0):
                tokenValido(tokenok)
            else:
                error = mb.showerror('Error', 'Falta ingresar datos requeridos')

        def ingresoNombre():
            extension = '.docx'
            futuro = fd.asksaveasfilename(defaultextension = extension,
                filetypes = [("docx", extension)], confirmoverwrite=True)
            if len(futuro) > 0:
                if not futuro.endswith(extension):
                    futuro += extension
                generarInforme(futuro)

        def obtenerMes(numero):
            mes = ('','enero','febrero',
                'marzo','abril','mayo',
                'junio','julio','agosto',
                'septiembre','octubre',
                'noviembre','diciembre')
            return mes[numero]

        def generarInforme(futuro):

            def getfanpageObject():
                objetoPage = graph.request('me?fields=id,name')
                return objetoPage

            def getinstaID(pageid):
                objetoPage = graph.request('/' + pageid + 
                    '?fields=instagram_business_account')
                instaid = objetoPage['instagram_business_account']['id']
                return instaid

            def replaceTime(key):
                temp = key['timestamp'].replace("T", " ")
                temp, _ = temp.split('+')
                temp = dt.datetime.strptime(temp, "%Y-%m-%d %H:%M:%S")
                return temp

            def acordePeriodo(identificadores, inicio, fin):
                id_en_periodo = []
                for i in identificadores:
                    post = graph.request('/'+ i + '?fields=timestamp')
                    post['timestamp'] = replaceTime(post)
                    if post['timestamp'] >= inicio and post['timestamp'] <= fin:
                        id_en_periodo.append(i)    
                return id_en_periodo

            def todoMedia(id_or_len):
                #Extrae todos los posts en el periodo indicado
                identificadores = []
                todoslosposts = graph.request('/'+ instaid +'/media')
                datostodos = todoslosposts['data']
                for x in datostodos:
                    identificadores.append(x['id'])
                if id_or_len is True:
                    return identificadores
                else:
                    return len(identificadores)

            #Extrae valor de la métrica de cada uno de los posts del periodo seleccionado
            def metricaUnica(metrica, identificadores, una_o_dos):
                top = []
                top_metrica = []
                for i in identificadores:
                    metricapost = graph.request('/'+ i +'/insights?metric='+ metrica)
                    datospost = metricapost['data']
                    for x in datospost:
                        u = x['values'][0]['value']
                        top.append(u)
                        top_metrica.append((i, u))
                if una_o_dos is True:
                    return top, top_metrica
                else:
                    return top

            def metricaAgregada(valores):
                if valores != 0:
                    return (sum(valores)/len(valores))
                else:
                    return 0

            def metricaSumada(valores):
                return sum(valores)

            def metricasPostPopular(postpopular):
                dicmetpop = {}
                if postpopular['media_type'] == 'IMAGE':
                    popmetricas = 'impressions,reach,engagement,saved'
                elif postpopular['media_type'] == 'VIDEO':
                    popmetricas = 'impressions,reach,engagement,saved,video_views'
                poppostmetricas = graph.request('/' + postpopular['id'] + '/insights?metric=' + popmetricas)
                accesodata = poppostmetricas['data']
                dicmetpop['impresiones'] = accesodata[0]['values'][0]['value']
                dicmetpop['alcance'] = accesodata[1]['values'][0]['value']
                dicmetpop['engagement'] = accesodata[2]['values'][0]['value']
                dicmetpop['guardado'] = accesodata[3]['values'][0]['value']
                if len(accesodata) == 5:
                    dicmetpop['vistas_video'] = accesodata[4]['values'][0]['value']

                return dicmetpop

            #TODO Extraer post popular a partir del engagement rate
            def postPopular(metrica, identificadores):

                top, top_metrica = metricaUnica(metrica, identificadores, True)

                #Extrae información del post más popular
                #TODO cachar cómo unpack una quinta variable, video views
                for a, b in top_metrica:
                    if b == max(top):
                        postpopular = graph.request('/'+ a + '?fields=id,caption,media_type,media_url,owner,timestamp,like_count,permalink')

                return postpopular

            tokenok = token.get()
            base = criterio.get()

            #Objeto con acceso a Facebook e Instagram
            graph = facebook.GraphAPI(access_token = tokenok, version = 3.0)

            #Geociencias PUC
            pageobjeto = getfanpageObject()
            pagename = pageobjeto['name']
            pageid = pageobjeto['id']
            instaid = getinstaID(pageid)

            #Rango temporal 2 para testear
            diainicio = int(dia_inicio.get())
            mesinicio = mesNumero(mes_inicio.get())
            anhoinicio = int(anho_inicio.get())
            diafin = int(dia_fin.get())
            mesfin = mesNumero(mes_fin.get())
            anhofin = int(anho_fin.get())
            inicito = dt.datetime(anhoinicio, mesinicio, diainicio)
            finito = dt.datetime(anhofin, mesfin, diafin)

            #Conversión a Timestamp string
            #inicio2 = str(int(time.mktime(inicito.timetuple()))) 
            #fin2 = str(int(time.mktime(finito.timetuple())))

            #Definir periodo analizado en el informe
            if int(inicito.year) == int(finito.year):
                if int(inicito.month) == int(finito.month):
                    mes = obtenerMes(inicito.month)
                else:
                    mes = (str(int(inicito.day)) + "/" + str(int(inicito.month)) + 
                        " al " + str(int(finito.day)) + "/" 
                        + str(int(finito.month)) + " del ")
                anho = int(finito.year)
            else:
                mes = (str(int(inicito.day)) + "/" + str(int(inicito.month)) + 
                    "/" + str(int(inicito.year))+ " al " + str(int(finito.day)) +
                     "/" + str(int(finito.month)) + "/" + str(int(finito.year)))
                anho = ""

            #Nombres de los archivos guardados
            desktop = os.getcwd()
            screenshot = os.path.join(desktop, 'screenshotfacebookpost.png')

            #Obtención de datos
            identificadores = todoMedia(True)
            identificadores = acordePeriodo(identificadores, inicito, finito)

            totaldeposts = len(identificadores)

            valoresimpresiones = metricaUnica('impressions', identificadores, False)
            impresionespromedio = metricaAgregada(valoresimpresiones)
            valoresreach = metricaUnica('reach', identificadores, False)
            reachpromedio = metricaAgregada(valoresreach)
            valoressaved = metricaUnica('saved', identificadores, False)
            savedtotal = metricaSumada(valoressaved)

            infopostpop = postPopular(base, identificadores)
            diccionariopop = metricasPostPopular(infopostpop)

            impresionespop = diccionariopop['impresiones']
            engagementpop = diccionariopop['engagement']
            reachpop = diccionariopop['alcance']
            savedpop = diccionariopop['guardado']
            if 'vistas_video' in diccionariopop:
                vistaspop = diccionariopop['vistas_video']
            
            linkpostinsta = infopostpop['permalink']

            #Edición de datos para display en documento
            impresionespromedio = str(round(impresionespromedio, 1)).replace('.',',')
            reachpromedio = str(round(reachpromedio, 1)).replace('.',',')

            #Escritura del documento
            document = Document()
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('Informe {0} {1} {2}'.format(pagename, mes, anho)).bold = True

            #Cuenta de Instagram
            pgram = document.add_paragraph()
            pgram.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pgram.add_run('Instagram').bold = True

            pgram1 = document.add_paragraph()
            pgram1.add_run('{0} posts realizados.'.format(totaldeposts))

            pgram2 = document.add_paragraph()
            pgram2.add_run('{0} impresiones (en promedio).'.format(impresionespromedio))

            pgram3 = document.add_paragraph()
            pgram3.add_run('{0} usuarios alcanzados (en promedio).'.format(reachpromedio))

            pgram4 = document.add_paragraph()
            pgram4.add_run('{0} posts guardados.'.format(savedtotal))

            pgram5 = document.add_paragraph('Post destacado:\n')

            #Extracción post más popular usando servicio APILEAP
            try:
                params = urlencode({'url':'{0}'.format(linkpostinsta), 
                    'access_key':'3dfe7afa8bafb96cdb0215bb12974a89'})
                urlretrieve('https://api.screenshotlayer.com/api/capture?' 
                    + params, screenshot)
                pimage = document.add_picture(screenshot, width=Inches(5))
            except:
                pass

            pgram6 = document.add_paragraph()
            pgram6.add_run('{0} interacciones de usuarios (likes + comentarios).'.format(engagementpop))

            pgram7 = document.add_paragraph()
            pgram7.add_run('{0} impresiones.'.format(impresionespop))

            pgram8 = document.add_paragraph()
            pgram8.add_run('{0} usuarios alcanzados.'.format(reachpop))

            document.add_page_break()

            #Guardado del documento
            document.save('{0}'.format(futuro))

            #Eliminación de imágenes
            if os.path.isfile(screenshot):
                os.remove(screenshot)

            #Apertura de archivo
            if sys.platform == 'linux':
                os.system("xdg-open " + futuro)
            elif sys.platform == 'win32':
                os.startfile(futuro)
            else:
                opener = "open" if sys.platform == 'darwin' else 'xdg-open'
                subprocess.call([opener, futuro])

        def habilitarMesInicio(event):
            listamesesi.config(state=['readonly'])

        def habilitarAnhoInicio(event):
            listaanhosi.config(state=['readonly'])

        def habilitarDiaFin(event):
            listadiasf.config(state=['readonly'])

        def habilitarMesFin(event):
            listamesesf.config(state=['readonly'])

        def habilitarAnhoFin(event):
            listaanhosf.config(state=['readonly'])

        def mesNumero(mestexto):
            textonumero = {'enero':1,'febrero':2,
                'marzo':3,'abril':4,'mayo':5,
                'junio':6,'julio':7,'agosto':8,
                'septiembre':9,'octubre':10,
                'noviembre':11,'diciembre':12}
            return textonumero[mestexto]

        def chequearTreinta(event):
            diainicio = int(dia_inicio.get())
            mesinicio = mesNumero(mes_inicio.get())
            anhoinicio = int(anho_inicio.get())
            diafin = int(dia_fin.get())
            mesfin = mesNumero(mes_fin.get())
            anhofin = int(anho_fin.get())
            fechainicio = dt.datetime(anhoinicio, mesinicio, diainicio)
            fechatermino = dt.datetime(anhofin, mesfin, diafin)
            if fechainicio < fechatermino:
                rangoidealdetermino = fechainicio + dt.timedelta(days=30)
                if fechatermino <= rangoidealdetermino:
                    pass
                else:
                    error = mb.showerror('Error', 'El rango entre fechas no debe superar los 30 días.')
                    listadiasi.current(0)
                    listamesesi.current(0)
                    listamesesi.state(['readonly','disabled'])
                    listaanhosi.current(0)
                    listaanhosi.state(['readonly','disabled'])
                    listadiasf.current(0)
                    listadiasf.state(['readonly','disabled'])
                    listamesesf.current(0)
                    listamesesf.state(['readonly','disabled'])
                    listaanhosf.current(0)
                    listaanhosf.state(['readonly','disabled'])
            else:
                error = mb.showerror('Error', 'La fecha de inicio debe ser anterior a la fecha de termino.')


        self.parent.columnconfigure(0, weight = 1)
        self.parent.rowconfigure(0, weight = 1)

        #Año actual para ingresar a la lista
        anhoactual = dt.datetime.today().year

        #Tupla con los meses en string
        mesestexto = ('enero','febrero',
                'marzo','abril','mayo',
                'junio','julio','agosto',
                'septiembre','octubre',
                'noviembre','diciembre')
        
        #Creación de base de datos
        desktop = os.getcwd()

        frame1 = ttk.Frame(self.parent, borderwidth = 1, 
            padding = (4, 4, 4, 4))
        frame1.grid(column = 0, row = 0, sticky = (W, E, S, N))

        frase = ttk.Label(frame1, text = 'Fecha Inicio: ')
        frase.grid(column = 0, row = 1)

        dia_inicio = StringVar()
        listadiasi = ttk.Combobox(frame1, textvariable = dia_inicio)
        listadiasi['values'] = [i for i in range(1, 32)]
        listadiasi.state(['readonly'])
        listadiasi.grid(column = 1, row = 1)
        listadiasi.bind("<<ComboboxSelected>>", habilitarMesInicio)

        mes_inicio = StringVar()
        listamesesi = ttk.Combobox(frame1, textvariable = mes_inicio)
        listamesesi['values'] = mesestexto
        listamesesi.state(['readonly', 'disabled'])
        listamesesi.grid(column = 2, row = 1)
        listamesesi.bind("<<ComboboxSelected>>", habilitarAnhoInicio)

        anho_inicio = StringVar()
        listaanhosi = ttk.Combobox(frame1, textvariable = anho_inicio)
        listaanhosi['values'] = [i for i in range(2010, anhoactual + 1)]
        listaanhosi.state(['readonly','disabled'])
        listaanhosi.grid(column = 3, row = 1)
        listaanhosi.bind("<<ComboboxSelected>>", habilitarDiaFin)

        frase2 = ttk.Label(frame1, text = 'Fecha Termino: ')
        frase2.grid(column = 0, row = 2)

        dia_fin = StringVar()
        listadiasf = ttk.Combobox(frame1, textvariable = dia_fin)
        listadiasf['values'] = [i for i in range(1, 32)]
        listadiasf.state(['readonly', 'disabled'])
        listadiasf.grid(column = 1, row = 2)
        listadiasf.bind("<<ComboboxSelected>>", habilitarMesFin)

        mes_fin = StringVar()
        listamesesf = ttk.Combobox(frame1, textvariable = mes_fin)
        listamesesf['values'] = mesestexto
        listamesesf.state(['readonly', 'disabled'])
        listamesesf.grid(column = 2, row = 2)
        listamesesf.bind("<<ComboboxSelected>>", habilitarAnhoFin)

        anho_fin = StringVar()
        listaanhosf = ttk.Combobox(frame1, textvariable = anho_fin)
        listaanhosf['values'] = [i for i in range(2010, anhoactual + 1)]
        listaanhosf.state(['readonly', 'disabled'])
        listaanhosf.grid(column = 3, row = 2)
        listaanhosf.bind("<<ComboboxSelected>>", chequearTreinta)

        s1 = ttk.Separator(frame1, orient = HORIZONTAL)
        s1.grid(row = 3, columnspan = 7, sticky = (W, E))

        seleccionPop = ttk.Label(frame1, 
            text = 'Seleccionar criterio para\n clasificar posts')
        seleccionPop.grid(column = 0, row = 4, sticky = (N), pady = 2)

        criterio = StringVar(None, '')
        por_impresion = ttk.Radiobutton(frame1, text = 'Impresiones', 
            variable = criterio, value = 'impressions')
        por_alcance = ttk.Radiobutton(frame1, text = 'Alcance', 
            variable = criterio, value = 'reach')
        #por_engagement = ttk.Radiobutton(frame1, text = 'Tasa de participación', 
        #    variable = criterio, value = 'engagement_rate')
        por_usuarios = ttk.Radiobutton(frame1, text = 'Acciones de usuarios', 
            variable = criterio, value = 'engagement')
        
        por_impresion.grid(column = 0, row = 5, sticky = W)
        por_alcance.grid(column = 0, row = 6, sticky = W)
        #por_engagement.grid(column = 0, row = 8, sticky = W)
        por_usuarios.grid(column = 0, row = 7, sticky = W)

        s2 = ttk.Separator(frame1, orient=VERTICAL)
        s2.grid(column = 1, row = 3, rowspan = 5, sticky = (N, S))

        ingreso_archivos = ttk.Label(frame1, text = 'Ingrese token')
        ingreso_archivos.grid(column = 2, row = 4, sticky = (E, W), 
            columnspan = 2, pady = 2)

        token = StringVar()
        entrada_post = ttk.Entry(frame1, textvariable = token)
        entrada_post.grid(column = 2, row = 5, sticky = (E, W))

        s3 = ttk.Separator(frame1, orient = HORIZONTAL)
        s3.grid(row = 9, column = 0, columnspan = 8, sticky = (E, W))

        generar = ttk.Button(frame1, text = 'Generar Informe', 
            command = chequearGenerar)
        generar.grid(row = 10, column = 0, columnspan = 5)

def main():
    root = Tk()
    root.title('Informe de Instagram')
    app = programo(root)
    root.mainloop()

if __name__ == "__main__":
    main()
