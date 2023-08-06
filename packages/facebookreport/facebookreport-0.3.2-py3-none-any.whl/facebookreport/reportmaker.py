#!/usr/bin/python
# -*- coding: utf-8 -*-

import pandas as pd
import datetime as dt
from statistics import mean
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from urllib.parse import urlencode
from urllib.request import urlretrieve
import sys
from tkinter import *
from tkinter import ttk
from tkinter import filedialog as fd
from tkinter import messagebox as mb
import pickle
import os
import subprocess
if sys.platform == 'win32':
    sys._enablelegacywindowsfsencoding()
if sys.platform == 'darwin':
    import matplotlib
    matplotlib.use('TkAgg')
import matplotlib.pyplot as plt

class programo(Tk):
    def __init__(self, parent):
        self.parent = parent
        self.iniciar()

    def iniciar(self):

        def chequearArchivo(tipo):
            filepath = filepage.get() if tipo == 'page' else filepost.get()
            fb = pd.read_csv(filepath, sep=",", skiprows=[1], encoding="latin-1")
            if ((tipo == 'post' and 'Lifetime Post Total Reach' in fb.columns) 
                or (tipo == 'page' and 'Weekly Page Engaged Users' in fb.columns)):
                pass
            else:
                error = mb.showerror('Error', 
                    'El archivo ingresado no corresponde al archivo requerido.')
                if tipo == 'post':
                    filepost.set('')
                else: 
                    filepage.set('')

        def chequearGenerar():
            filepostpath = filepost.get()
            filepagepath = filepage.get()
            listaopcion = lista.get()
            if (len(filepostpath) != 0
                and len(filepagepath) != 0
                and os.path.exists(filepostpath)
                and os.path.exists(filepagepath)
                and lista.current() != -1
                and len(criterio.get()) > 0):
                ingresoNombre()
            else:
                error = mb.showerror('Error', 'Falta ingresar datos requeridos')

        def ingresoNombre():
            extension = '.docx'
            futuro = fd.asksaveasfilename(defaultextension = extension,
                filetypes = [("docx", extension)], confirmoverwrite=True)
            if len(futuro) > 0:
                if not futuro.endswith(extension):
                    futuro += extension
                generarInforme(futuro)

        def getDataFrame(path, date):
            dataframe = pd.read_csv(path, sep=",", 
                skiprows=[1], encoding="latin-1")
            dataframe.columns = dataframe.columns.str.replace('\s', '_')
            dataframe[date] = pd.to_datetime(dataframe[date])
            dataframe = dataframe.fillna(0)
            return dataframe

        def obtenerMes(numero):
            mes = ('','enero','febrero',
                'marzo','abril','mayo',
                'junio','julio','agosto',
                'septiembre','octubre',
                'noviembre','diciembre')
            return mes[numero]

        def generarInforme(futuro):

            cuentaName = cuenta.get()
            filepathpost = filepost.get()
            filepathpage = filepage.get()
            fechapost = 'Posted'
            fechapage = 'Date'

            #Ingreso y edición del archivo csv con estadísticas de la Facebook fanpage
            fb = getDataFrame(filepathpost, fechapost)
            fb2 = getDataFrame(filepathpage, fechapage)

            totalposts = len(fb.index)

            fb['engagement_rate'] = (fb['Lifetime_Engaged_Users'] 
                / fb['Lifetime_Post_Total_Reach'])
            fb2['engagement_rate'] = (fb2['Daily_Page_Engaged_Users'] 
                / fb2['Daily_Total_Reach'])

            base = criterio.get()

            #Sacar el promedio de engagement rate
            mediaER = mean(fb['engagement_rate']) #A nivel post
            mediaERpag = mean(fb2['engagement_rate']) #A nivel page

            #Extracción de información del post más popular
            pd.options.display.max_colwidth = 100
            pop = fb.nlargest(1, base)
            maslink = pop['Permalink']
            mastime = pop[fechapost]
            mastype = pop['Type']
            masreach = int(pop['Lifetime_Post_Total_Reach'])
            masimp = int(pop['Lifetime_Post_Total_Impressions'])
            maseng = int(pop['Lifetime_Engaged_Users'])
            masshared = int(pop['Lifetime_Post_Stories_by_action_type_-_share'])
            masER = float(pop['engagement_rate'])

            #Extracción de informacion a nivel página
            likesAhora = fb2['Lifetime_Total_Likes'].iloc[-1]
            newlikes = (fb2['Lifetime_Total_Likes'].iloc[-1] - 
                fb2['Lifetime_Total_Likes'].iloc[0])
            usuarios = mean(fb2['Daily_Page_Engaged_Users'])
            alcance = mean(fb2['Daily_Total_Reach'])
            impresiones = mean(fb2['Daily_Total_Impressions'])

            #Edición de datos para display en documento
            usuarios = str(round(usuarios, 1)).replace('.',',')
            alcance = str(round(alcance, 1)).replace('.',',')
            impresiones = str(round(impresiones, 1)).replace('.',',')
            masER = str(round(masER * 100, 1)).replace('.',',')
            mediaER = str(round(mediaER * 100, 1)).replace('.',',')

            #Definir periodo analizado en el informe
            primerDia = fb2[fechapage].iloc[0]
            ultimoDia = fb2[fechapage].iloc[-1]
            if int(primerDia.year) == int(ultimoDia.year):
                if int(primerDia.month) == int(ultimoDia.month):
                    mes = obtenerMes(primerDia.month)
                else:
                    mes = (str(int(primerDia.day)) + "/" + str(int(primerDia.month)) + 
                        " al " + str(int(ultimoDia.day)) + "/" 
                        + str(int(ultimoDia.month)) + " del ")
                anho = int(mastime.dt.year)
            else:
                mes = (str(int(primerDia.day)) + "/" + str(int(primerDia.month)) + 
                    "/" + str(int(primerDia.year))+ " al " + str(int(ultimoDia.day)) +
                     "/" + str(int(ultimoDia.month)) + "/" + str(int(ultimoDia.year)))
                anho = ""

            #Nombres de los archivos guardados
            desktop = os.getcwd()
            screenshot = os.path.join(desktop, 'screenshotfacebookpost.png')
            pnghora = os.path.join(desktop, 'mejorhora.png')
            pnglikes = os.path.join(desktop, 'totallikes.png')
            pngusers = os.path.join(desktop, 'engagedusers.png')
            pngreach = os.path.join(desktop, 'totalreach.png')
            pngimpres = os.path.join(desktop, 'totalimpressions.png')
            pngengage = os.path.join(desktop, 'totalengagement.png')

            #Top Ten más populares / Análisis de hora.
            pop2 = fb.nlargest(10, base)
            time = pop2[fechapost]
            fighora = time.groupby(pop2[fechapost].dt.hour).count().plot(kind='bar')
            plt.title('Horario de los mejores posts')
            plt.savefig(pnghora)
            plt.close()

            #Otros gráficos
            fig1 = fb2.plot(x='Date', y='Lifetime_Total_Likes')
            plt.title('Total de "me gusta" en el tiempo')
            plt.savefig(pnglikes)
            plt.close()
            fig2 = fb2.plot(x='Date', y='Daily_Page_Engaged_Users')
            plt.title('Usuarios que interactuaron en el tiempo')
            plt.savefig(pngusers)
            plt.close()
            fig3 = fb2.plot(x='Date', y='Daily_Total_Reach')
            plt.title('Alcance total en el tiempo')
            plt.savefig(pngreach)
            plt.close()
            fig4 = fb2.plot(x='Date', y='Daily_Total_Impressions')
            plt.title('Impresiones totales en el tiempo')
            plt.savefig(pngimpres)
            plt.close()
            fig5 = fb2.plot(x='Date', y='engagement_rate')
            plt.title('Tasa de participación en el tiempo')
            plt.savefig(pngengage)
            plt.close()

            #Escritura del documento
            document = Document()
            p = document.add_paragraph()
            p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run('Informe {0} {1} {2}'.format(cuentaName, mes, anho)).bold = True

            #Cuenta de Facebook
            p2 = document.add_paragraph()
            p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.add_run('Facebook').bold = True

            p3 = document.add_paragraph()
            p3.add_run('{0} nuevos likes (llegando a {1}).'.format(newlikes, likesAhora))

            p4 = document.add_paragraph()
            p4.add_run('{0} posts emitidos.'.format(totalposts))

            p5 = document.add_paragraph()
            p5.add_run('{0} usuarios en promedio que realizaron algún tipo de interacción.'.format(usuarios))

            p6 = document.add_paragraph()
            p6.add_run('{0} personas alcanzadas (promedio de usuarios únicos en cada post).'.format(alcance))

            p7 = document.add_paragraph()
            p7.add_run('{0} impresiones en promedio.'.format(impresiones))

            p75 = document.add_paragraph()
            p75.add_run('{0}% tasa de participación (engagement rate) promediada.'.format(mediaER))

            p8 = document.add_paragraph('Post destacado:\n')
            #Porque maslink es un pandas Series
            linkpost = str(maslink.values[0])
            p8.add_run('{0}'.format(linkpost))

            #Acceso a internet y extración de imagen del post más popular.
            try:
                params = urlencode({'url':'{0}'.format(linkpost), 
                    'access_key':'3dfe7afa8bafb96cdb0215bb12974a89'})
                urlretrieve('https://api.screenshotlayer.com/api/capture?' 
                    + params, screenshot)
                pimage = document.add_picture(screenshot, width=Inches(5))
            except:
                pass

            p85 = document.add_paragraph()
            p85.add_run('{0}% tasa de participación.'.format(masER))

            p9 = document.add_paragraph()
            p9.add_run('{0} usuarios únicos alcanzados.'.format(masreach))

            p10 = document.add_paragraph()
            p10.add_run('{0} impresiones.'.format(masimp))

            p11 = document.add_paragraph()
            p11.add_run('{0} interacciones.'.format(maseng))

            if int(masshared) > 0:
                p12 = document.add_paragraph()
                p12.add_run('{0} veces compartido.'.format(masshared))

            document.add_page_break()

            #Sección de gráficos
            pimage2 = document.add_picture(pnglikes, width=Inches(5))
            pimage3 = document.add_picture(pngimpres, width=Inches(5))

            pimage4 = document.add_picture(pngusers, width=Inches(5))
            pimage5 = document.add_picture(pngreach, width=Inches(5))

            pimage6 = document.add_picture(pngengage, width=Inches(5))
            pimage7 = document.add_picture(pnghora, width=Inches(5))

            document.add_page_break()

            #Guardado del documento
            document.save('{0}'.format(futuro))

            #Eliminación de imágenes
            os.remove(pnglikes)
            os.remove(pnghora)
            os.remove(pngengage)
            os.remove(pngreach)
            os.remove(pngusers)
            os.remove(pngimpres)
            if os.path.isfile(screenshot):
                os.remove(screenshot)

            #Apertura de archivo
            if sys.platform == 'linux':
                os.system("xdg-open " + futuro)
            elif sys.platform == 'win32':
                os.startfile(futuro)
            else:
                opener = "open" if sys.platform == 'darwin' else 'xdg-open'
                subprocess.call([opener, futuro])

        def seleccionArchivo(nivel):
            filename = fd.askopenfilename(defaultextension = '.csv', 
                filetypes = [('csv', '*.csv')])
            if filename:
                filepost.set(filename) if nivel is 'post' else filepage.set(filename)
                chequearArchivo(nivel)

        def confirmarEliminar(nombre):
            confirmar = mb.askokcancel('Eliminar cuenta', 
                '¿Estás seguro que quieres eliminar esta cuenta?')
            if confirmar == True:
                lista_cuenta.remove(nombre)
                escribirBD = open(baseDatos, 'wb')
                pickle.dump(lista_cuenta, escribirBD)
                escribirBD.close()
                lista.config(values = lista_cuenta)
                textoDone.set('OK')
                done.after(5000, lambda: textoDone.set(''))
                cuenta.set('')

        def eliminarCuenta():
            nombre = cuenta.get()
            if nombre in lista_cuenta:
                confirmarEliminar(nombre)

        def nuevaCuenta():
            nombre = nombre_cuenta.get()
            if len(nombre) > 2:
                lista_cuenta.append(nombre)
                escribirBD = open(baseDatos, 'wb')
                pickle.dump(lista_cuenta, escribirBD)
                escribirBD.close()
                lista.config(values = lista_cuenta)
                textoDone.set('OK')
                done.after(5000, lambda: textoDone.set(''))
                toplevel.destroy()
                lista.set(nombre)
            else:
                error = mb.showerror('Error', 
                    'Ingresa un nombre válido (mayor a 2 carácteres)')

        def dialogoCuenta():
            global toplevel 
            toplevel = Toplevel()
            toplevel.title("Crear nueva cuenta")
            mensaje = ttk.Label(toplevel, text = "Escriba nombre de la cuenta")
            mensaje.grid(column = 0, row = 0, sticky = (W, E))
            entrada = ttk.Entry(toplevel, textvariable = nombre_cuenta)
            entrada.grid(column = 0, row = 1, sticky = (W, E))
            boton = ttk.Button(toplevel, text="Guardar", command = nuevaCuenta)
            boton.grid(column = 1, row = 1)


        self.parent.columnconfigure(0, weight = 1)
        self.parent.rowconfigure(0, weight = 1)

        #Creación de base de datos
        desktop = os.getcwd()
        baseDatos = os.path.join(desktop, 'cuentas.fbr')

        if os.path.isfile('cuentas.fbr'):
            leerBD = open(baseDatos, 'rb')
            lista_cuenta = pickle.load(leerBD)
            leerBD.close()
        else:
            lista_cuenta = []

        frame1 = ttk.Frame(self.parent, borderwidth = 1, 
            padding = (4, 4, 4, 4))
        frame1.grid(column = 0, row = 0, sticky = (W, E, S, N))

        frase = ttk.Label(frame1, text = 'Cuenta: ')
        frase.grid(column = 0, row = 1)

        nombre_cuenta = StringVar()
        cuenta = StringVar()
        lista = ttk.Combobox(frame1, textvariable = cuenta)
        lista['values'] = (lista_cuenta)
        lista.state(['readonly'])
        lista.grid(column = 1, row = 1)
        if len(lista_cuenta) > 0:
            lista.current(0)

        agregar_cuenta = ttk.Button(frame1, 
            text = 'Nueva cuenta', command = dialogoCuenta)
        agregar_cuenta.grid(column = 2, row = 1, padx = 2)

        eliminar_cuenta = ttk.Button(frame1, 
            text = 'Eliminar cuenta', command = eliminarCuenta)
        eliminar_cuenta.grid(column = 3, row = 1)

        textoDone = StringVar()
        done = ttk.Label(frame1, textvariable = textoDone)
        done.grid(column = 4, row = 1)

        s1 = ttk.Separator(frame1, orient = HORIZONTAL)
        s1.grid(row = 2, columnspan = 7, sticky = (W, E))

        seleccionPop = ttk.Label(frame1, 
            text = 'Seleccionar criterio para\n clasificar posts')
        seleccionPop.grid(column = 0, row = 3, sticky = (N), pady = 2)

        criterio = StringVar(None, '')
        por_impresion = ttk.Radiobutton(frame1, text = 'Impresiones', 
            variable = criterio, value = 'Lifetime_Post_Total_Impressions')
        por_alcance = ttk.Radiobutton(frame1, text = 'Alcance', 
            variable = criterio, value = 'Lifetime_Post_Total_Reach')
        por_engagement = ttk.Radiobutton(frame1, text = 'Tasa de participación', 
            variable = criterio, value = 'engagement_rate')
        por_usuarios = ttk.Radiobutton(frame1, text = 'Usuarios involucrados', 
            variable = criterio, value = 'Lifetime_Engaged_Users')
        
        por_impresion.grid(column = 0, row = 4, sticky = W)
        por_alcance.grid(column = 0, row = 5, sticky = W)
        por_engagement.grid(column = 0, row = 6, sticky = W)
        por_usuarios.grid(column = 0, row = 7, sticky = W)

        s2 = ttk.Separator(frame1, orient=VERTICAL)
        s2.grid(column = 1, row = 3, rowspan = 5, sticky = (N, S))

        ingreso_archivos = ttk.Label(frame1, text = 'Seleccione archivos csv')
        ingreso_archivos.grid(column = 2, row = 3, sticky = (E, W), 
            columnspan = 2, pady = 2)

        nivel_post = ttk.Label(frame1, text = 'Nivel post')
        nivel_post.grid(column = 2, row = 4)

        nivel_page = ttk.Label(frame1, text = 'Nivel page')
        nivel_page.grid(column = 2, row = 5)

        filepost = StringVar()
        entrada_post = ttk.Entry(frame1, textvariable = filepost)
        entrada_post.grid(column = 3, row = 4)

        filepage = StringVar()
        entrada_page = ttk.Entry(frame1, textvariable = filepage)
        entrada_page.grid(column = 3, row = 5)

        boton_post = ttk.Button(frame1, text = "..." , 
            command = lambda: seleccionArchivo('post'))
        boton_post.grid(column = 4, row = 4)
        boton_page = ttk.Button(frame1, text = "..." , 
            command = lambda: seleccionArchivo('page'))
        boton_page.grid(column = 4, row = 5)

        s3 = ttk.Separator(frame1, orient = HORIZONTAL)
        s3.grid(row = 8, column = 0, columnspan = 8, sticky = (E, W))

        generar = ttk.Button(frame1, text = 'Generar Informe', 
            command = chequearGenerar)
        generar.grid(row = 9, column = 0, columnspan = 5)

def main():
    root = Tk()
    root.title('Informe de Facebook')
    app = programo(root)
    root.mainloop()

if __name__ == "__main__":
    main()